# ================================================
# 2_📘_Weekly_NFR.py — ScopeSight v3.2 (Enhanced UI)
# Weekly NFR Consolidator (Project-based)
# ================================================

import io
import re
import os
import json
import datetime as dt
from typing import List, Dict, Tuple, Optional

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

# Core UI
from modules.ui_sidebar import render_sidebar
from modules.ui_hide_nav import hide_streamlit_nav
from modules.ui_branding import set_pmo_theme, pmo_footer

# DB
from modules.db import (
    run_query,
    insert_weekly_nfr,
    save_weekly_nfr_file,
)

# Logging
from modules.log_utils import log_event, log_document

# Security
from auth.login import require_login
require_login()

# Notifications (NEW system)
from modules.notifications_config import get_default_severity
from modules.notifications_utils import emit_event

# -------------------------------------------------
# Weekly consolidation agents (T1/T2)
# -------------------------------------------------
try:
    from modules.nfr.nfr_weekly_agent_T1 import consolidate_week_with_agent_t1
except Exception:
    consolidate_week_with_agent_t1 = None

try:
    from modules.nfr.nfr_weekly_agent_T2 import consolidate_week_with_agent_t2
except Exception:
    consolidate_week_with_agent_t2 = None


# -------------------------
# Page bootstrap
# -------------------------
hide_streamlit_nav()
set_pmo_theme(page_title="📘 Weekly NFR Consolidator")

st.markdown(
    """
<style>
header[data-testid="stHeader"] {
    height: 0px !important;
    visibility: hidden !important;
}

/* Enhanced card styling */
.nfr-card {
    background: white;
    border: 2px solid #4facfe;
    padding: 1.5rem;
    border-radius: 12px;
    margin: 1.5rem 0;
    box-shadow: 0 4px 12px rgba(79, 172, 254, 0.15);
}

.nfr-card h3 {
    color: #0077be;
    margin: 0 0 1rem 0;
    font-size: 1.3rem;
    font-weight: 600;
}

.nfr-card .info-row {
    background: #f0f9ff;
    padding: 0.75rem 1rem;
    margin: 0.5rem 0;
    border-radius: 6px;
    border-left: 4px solid #4facfe;
}

.nfr-card .info-row strong {
    color: #0077be;
}

/* Section headers */
.section-header {
    background: linear-gradient(90deg, #4facfe 0%, #00f2fe 100%);
    padding: 1rem 1.5rem;
    border-radius: 8px;
    margin: 2rem 0 1rem 0;
}

.section-header h3 {
    color: white;
    margin: 0;
    font-size: 1.2rem;
    font-weight: 600;
}

/* Step headers */
.step-header {
    background: #f0f9ff;
    border-left: 4px solid #4facfe;
    padding: 0.75rem 1rem;
    border-radius: 6px;
    margin: 1.5rem 0 1rem 0;
}

.step-header h4 {
    color: #0077be;
    margin: 0;
    font-size: 1.1rem;
    font-weight: 600;
}

/* Info boxes */
.info-box {
    background: #f0fff4;
    border-left: 4px solid #48bb78;
    padding: 1rem;
    border-radius: 4px;
    margin: 1rem 0;
}

/* Upload area styling */
.upload-section {
    background: #ffffff;
    border: 2px dashed #4facfe;
    border-radius: 8px;
    padding: 1rem;
    text-align: center;
    margin: 1rem 0;
}

/* Generate button enhancement */
div.stButton > button {
    background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
    color: white;
    font-size: 1.1rem;
    font-weight: 600;
    padding: 0.75rem 2rem;
    border: none;
    border-radius: 8px;
    transition: all 0.3s ease;
}

div.stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 12px rgba(79, 172, 254, 0.4);
}

/* Results card */
.results-card {
    background: white;
    border: 1px solid #e0e0e0;
    border-radius: 12px;
    padding: 2rem;
    margin: 2rem 0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.05);
}

/* Preview container */
.preview-container {
    background: white;
    border: 2px solid #4facfe;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1.5rem 0;
}
</style>
""",
    unsafe_allow_html=True,
)

render_sidebar()

# -------------------------
# Constants
# -------------------------
FONT_FAMILY = "Lexend Light"
FONT_SIZE = 11
SECTION_FONT_SIZE = 14
SECTION_GREEN = RGBColor(71, 163, 64)
DFMT = "%d %B %Y"
ACTIONS_HEADERS = ["Title", "Detail", "Owner", "Due Date"]


# -------------------------
# JSON helper
# -------------------------
def _parse_json(value):
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            return json.loads(value)
        except Exception:
            return {}
    return {}


# -------------------------
# Notifications helper (safe)
# -------------------------
def safe_emit_event(
    *,
    event_type: str,
    project_id: int,
    actor_email: str,
    title: str,
    body: str,
    entity_type: str,
    entity_id: Optional[int],
):
    """
    Never break the user flow if notifications fail.
    This will integrate with your v5 notifications page (user opt-ins handled downstream).
    """
    try:
        emit_event(
            event_type=event_type,
            project_id=int(project_id),
            actor_email=(actor_email or "").strip().lower(),
            title=title,
            body=body,
            entity_type=entity_type,
            entity_id=int(entity_id) if entity_id is not None else None,
            severity=get_default_severity(event_type),
        )
    except Exception:
        pass


# -------------------------
# Tier helper (safe)
# -------------------------
def get_project_tier(project_id: int, merged_config: dict) -> str:
    """
    Weekly page needs a tier to route the correct agent.
    We try to fetch a tier-like column from projects, but your schema
    may vary. If it fails, we fall back to client merged_config or "T2".
    """
    tier = None

    # Try common variants safely (schema differences)
    try:
        df = run_query("SELECT tier FROM projects WHERE project_id = :pid", {"pid": project_id})
        if df is not None and not df.empty:
            tier = df.iloc[0].get("tier")
    except Exception:
        tier = None

    if not tier:
        tier = merged_config.get("tier") or merged_config.get("project_tier") or "T2"

    tier = str(tier).strip().upper()
    if tier in ("TIER_1", "1", "T1"):
        return "T1"
    if tier in ("TIER_2", "2", "T2"):
        return "T2"
    return "T2"


# -------------------------
# Project selector (client_scaffold → projects)
# -------------------------
def select_project():
    """
    Select CLIENT from client_scaffold → PROJECT from projects table.

    NFR defaults can live in:
        • client_scaffold.settings -> ["nfr_defaults"]
        • client_scaffold.nfr_config (top-level)
    """
    clients = run_query(
        """
        SELECT 
            id           AS client_id,
            client_name,
            client_code,
            nfr_config,
            settings
        FROM client_scaffold
        WHERE status = 'approved'
        ORDER BY client_name
    """
    )

    if clients is None or clients.empty:
        st.error("⚠ No approved clients found. Create and approve a client first.")
        # always return the same shape
        return None, None, None, None, {}, "", ""

    client_name = st.selectbox("Select Client", clients["client_name"])
    c_row = clients.loc[clients["client_name"] == client_name].iloc[0]
    client_id = int(c_row["client_id"])
    client_code = c_row.get("client_code", "") or ""

    settings_obj = _parse_json(c_row.get("settings"))
    nfr_cfg_obj = _parse_json(c_row.get("nfr_config"))

    merged_config = {}
    if isinstance(settings_obj, dict):
        base_defaults = settings_obj.get("nfr_defaults", {})
        if isinstance(base_defaults, dict):
            merged_config.update(base_defaults)
    if isinstance(nfr_cfg_obj, dict):
        merged_config.update(nfr_cfg_obj)

    projects = run_query(
        """
        SELECT project_id, project_name, project_code
        FROM projects
        WHERE client_id = :cid
          AND LOWER(status) = 'open'
        ORDER BY project_name
    """,
        {"cid": client_id},
    )

    if projects is None or projects.empty:
        st.warning("⚠ This client has no open projects yet.")
        return None, None, client_name, client_id, merged_config, client_code, ""

    project_name = st.selectbox("Select Project", projects["project_name"])
    p_row = projects.loc[projects["project_name"] == project_name].iloc[0]
    project_id = int(p_row["project_id"])
    project_code = p_row.get("project_code", "") or ""

    return project_name, project_id, client_name, client_id, merged_config, client_code, project_code


# -------------------------
# Date helpers
# -------------------------
def monday_of(date_):
    return date_ - dt.timedelta(days=date_.weekday())


def friday_of(date_):
    return monday_of(date_) + dt.timedelta(days=4)


# -------------------------
# Word parser helpers
# -------------------------
def read_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join([p for p in parts if p])


def find_actions_tables(doc: Document):
    out = []
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        headers = [c.text.strip().lower() for c in tbl.rows[0].cells]
        if len(headers) < 4:
            continue

        if (
            ("title" in headers[0])
            and ("detail" in headers[1])
            and ("owner" in headers[2])
            and ("due" in headers[3] or "date" in headers[3])
        ):
            rows = []
            for r in tbl.rows[1:]:
                cells = [c.text.strip() for c in r.cells]
                if not any(cells):
                    continue
                rows.append(
                    {
                        "Title": cells[0],
                        "Detail": cells[1],
                        "Owner": cells[2],
                        "Due Date": cells[3],
                    }
                )
            out.append(rows)
    return out


def parse_due_date(text: str):
    text = (text or "").strip()
    if not text:
        return None
    for fmt in ["%d %B %Y", "%d/%m/%Y", "%Y-%m-%d", "%d.%m.%Y"]:
        try:
            return dt.datetime.strptime(text, fmt).date()
        except Exception:
            pass
    return None


def consolidate_actions(tables):
    def key(r):
        return (r.get("Title", "").lower(), r.get("Detail", "").lower())

    result = {}
    for tbl in tables:
        for row in tbl:
            if not any(row.values()):
                continue
            k = key(row)
            if k not in result:
                result[k] = row.copy()
            else:
                existing = result[k]

                if not existing.get("Owner") and row.get("Owner"):
                    existing["Owner"] = row["Owner"]

                d0 = parse_due_date(existing.get("Due Date"))
                d1 = parse_due_date(row.get("Due Date"))
                if d1 and (not d0 or d1 < d0):
                    existing["Due Date"] = row["Due Date"]

    return list(result.values())


# -------------------------
# DOCX builders
# -------------------------
def _apply_run_font(run, size_pt=FONT_SIZE, bold=False, colour=None):
    run.font.name = FONT_FAMILY
    r = run._element.rPr.rFonts
    r.set(qn("w:ascii"), FONT_FAMILY)
    r.set(qn("w:hAnsi"), FONT_FAMILY)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if colour:
        run.font.color.rgb = colour


def add_section_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _apply_run_font(run, size_pt=SECTION_FONT_SIZE, colour=SECTION_GREEN)
    return p


def add_blank_line(doc: Document):
    doc.add_paragraph("")


def build_overview_table(doc: Document, overview: Dict[str, str]):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Week commencing", "Date range", "Meeting title", "Project / Client"]
    for i, h in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        _apply_run_font(run, bold=True)

    vals = [
        overview.get("Week commencing", ""),
        overview.get("Date range", ""),
        overview.get("Meeting title", ""),
        overview.get("Project / Client", ""),
    ]

    row = table.add_row().cells
    for i, v in enumerate(vals):
        run = row[i].paragraphs[0].add_run(v)
        _apply_run_font(run)


def build_attendees_table(doc: Document, internal, external):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    headers = ["Internal", "External"]
    for i, h in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        _apply_run_font(run, bold=True)

    row = table.add_row().cells
    run_i = row[0].paragraphs[0].add_run("\n".join(internal))
    run_e = row[1].paragraphs[0].add_run("\n".join(external))
    _apply_run_font(run_i)
    _apply_run_font(run_e)


def build_discussion(doc: Document, sections):
    for subhead, bullets in sections:
        p = doc.add_paragraph()
        run = p.add_run(subhead)
        _apply_run_font(run, bold=True)

        for b in bullets:
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent = Inches(0.25)
            run = bp.add_run(f"• {b}")
            _apply_run_font(run)


def build_actions_table(doc: Document, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    for i, h in enumerate(ACTIONS_HEADERS):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        _apply_run_font(run, bold=True)

    if not rows:
        r = table.add_row().cells
        vals = ["No New Actions", "", "", ""]
        for i, v in enumerate(vals):
            run = r[i].paragraphs[0].add_run(v)
            _apply_run_font(run)
        return

    for row in rows:
        r = table.add_row().cells
        vals = [
            row.get("Title", ""),
            row.get("Detail", ""),
            row.get("Owner", ""),
            row.get("Due Date", ""),
        ]
        for i, v in enumerate(vals):
            run = r[i].paragraphs[0].add_run(v)
            _apply_run_font(run)


def make_weekly_docx(
    overview,
    objectives,
    attendees_int,
    attendees_ext,
    discussion,
    actions,
) -> bytes:
    doc = Document()

    add_section_title(doc, "Overview")
    build_overview_table(doc, overview)
    add_blank_line(doc)

    add_section_title(doc, "Objectives / Agenda")
    p = doc.add_paragraph()
    run = p.add_run(objectives)
    _apply_run_font(run)
    add_blank_line(doc)

    add_section_title(doc, "Attendees")
    build_attendees_table(doc, attendees_int, attendees_ext)
    add_blank_line(doc)

    add_section_title(doc, "Key Discussion Points")
    if discussion:
        build_discussion(doc, discussion)
    else:
        p = doc.add_paragraph()
        run = p.add_run("(No discussion points extracted.)")
        _apply_run_font(run)
    add_blank_line(doc)

    add_section_title(doc, "New Actions")
    build_actions_table(doc, actions)
    add_blank_line(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -------------------------
# HTML preview
# -------------------------
def as_html_preview(overview, objectives, attendees_int, attendees_ext, discussion, actions):
    css = f"""
    <style>
    body {{
        font-family: '{FONT_FAMILY}', Arial, sans-serif;
        font-size:{FONT_SIZE}pt;
    }}
    h2 {{
        color: rgb(71,163,64);
        font-size:{SECTION_FONT_SIZE}pt;
    }}
    table {{
        border-collapse: collapse;
        width: 100%;
        margin: 1rem 0;
    }}
    th, td {{
        border: 1px solid #ddd;
        padding: 8px;
        vertical-align: top;
    }}
    th {{
        font-weight: bold;
        background: #f0f9ff;
    }}
    </style>
    """

    ov_tbl = f"""
    <table>
        <tr>
            <th>Week commencing</th><th>Date range</th><th>Meeting title</th><th>Project / Client</th>
        </tr>
        <tr>
            <td>{overview.get("Week commencing","")}</td>
            <td>{overview.get("Date range","")}</td>
            <td>{overview.get("Meeting title","")}</td>
            <td>{overview.get("Project / Client","")}</td>
        </tr>
    </table>
    """

    att_tbl = f"""
    <table>
        <tr><th>Internal</th><th>External</th></tr>
        <tr>
            <td>{'<br/>'.join(attendees_int)}</td>
            <td>{'<br/>'.join(attendees_ext)}</td>
        </tr>
    </table>
    """

    disc_html = []
    for sub, bullets in discussion:
        disc_html.append(f"<div style='margin-top:1rem;'><b>{sub}</b></div>")
        if bullets:
            disc_html.append("<ul>" + "".join(f"<li>{b}</li>" for b in bullets) + "</ul>")
    if not disc_html:
        disc_html.append("<i>(No discussion points extracted)</i>")

    if actions:
        rows_html = "".join(
            f"<tr><td>{a.get('Title','')}</td><td>{a.get('Detail','')}</td>"
            f"<td>{a.get('Owner','')}</td><td>{a.get('Due Date','')}</td></tr>"
            for a in actions
        )
    else:
        rows_html = "<tr><td>No New Actions</td><td></td><td></td><td></td></tr>"

    act_tbl = f"""
    <table>
        <tr><th>Title</th><th>Detail</th><th>Owner</th><th>Due Date</th></tr>
        {rows_html}
    </table>
    """

    disc_joined = "".join(disc_html)

    return f"""
    {css}
    <h2>Overview</h2>{ov_tbl}
    <h2>Objectives / Agenda</h2><p>{objectives}</p>
    <h2>Attendees</h2>{att_tbl}
    <h2>Key Discussion Points</h2>{disc_joined}
    <h2>New Actions</h2>{act_tbl}
    """


# -------------------------
# Main UI
# -------------------------
def ui():
    st.markdown(
        """
    <div class='info-box'>
        <strong style='color: #48bb78;'>💡 Recommended Workflow</strong><br/>
        <span style='color: #2d3748;'>Generate Daily NFRs throughout the week first, then consolidate them here into a comprehensive Weekly NFR.</span>
    </div>
    """,
        unsafe_allow_html=True,
    )

    # Project selector
    st.markdown(
        """
    <div class='section-header'>
        <h3>📁 Project Selection</h3>
    </div>
    """,
        unsafe_allow_html=True,
    )

    result = select_project()
    if not result or result[0] is None:
        st.stop()

    project_name, project_id, client_name, client_id, merged_config, client_code, project_code = result

    # Tier routing
    project_tier = get_project_tier(project_id, merged_config)

    # Week configuration
    st.markdown(
        """
    <div class='section-header'>
        <h3>📅 Week Configuration</h3>
    </div>
    """,
        unsafe_allow_html=True,
    )

    colA, colB = st.columns([2, 1])
    with colA:
        week_date = st.date_input("Week commencing (Monday)", value=monday_of(dt.date.today()))
        mon = monday_of(week_date)
        fri = friday_of(week_date)
        date_range_label = f"{mon.strftime(DFMT)} – {fri.strftime(DFMT)}"

    with colB:
        default_time = merged_config.get("default_time", "10:00")
        st.text_input("Default meeting time", value=default_time, key="default_time")

    st.info(f"📆 Week Range: **{date_range_label}**")

    # Input sections
    st.markdown(
        """
    <div class='section-header'>
        <h3>📝 Weekly NFR Content</h3>
    </div>
    """,
        unsafe_allow_html=True,
    )

    # Step 1: Monday transcript
    st.markdown(
        """
    <div class='step-header'>
        <h4>1️⃣ Monday Meeting Transcript</h4>
    </div>
    """,
        unsafe_allow_html=True,
    )

    c1, c2 = st.columns(2)

    monday_text = ""
    with c1:
        monday_text = st.text_area(
            "✍️ Paste Monday transcript",
            height=240,
            placeholder="Paste the Monday meeting transcript here...",
        )

    with c2:
        monday_docx = st.file_uploader("📎 Or upload Monday .docx", type=["docx"])
        if monday_docx and not monday_text:
            monday_text = read_docx_text(monday_docx.read())
            st.success(f"✅ Loaded: {monday_docx.name}")

    # Step 2: Daily NFRs
    st.markdown(
        """
    <div class='step-header'>
        <h4>2️⃣ Daily NFR Documents (Optional)</h4>
    </div>
    """,
        unsafe_allow_html=True,
    )

    weekly_files = st.file_uploader(
        "📤 Upload Daily NFR .docx files from the week",
        type=["docx"],
        accept_multiple_files=True,
    )

    if weekly_files:
        st.success(f"✅ {len(weekly_files)} file(s) uploaded")

    # Step 3: Attendees
    st.markdown(
        """
    <div class='step-header'>
        <h4>3️⃣ Meeting Attendees (Optional)</h4>
    </div>
    """,
        unsafe_allow_html=True,
    )

    a1, a2 = st.columns(2)
    with a1:
        att_int = st.text_area(
            "👥 Internal attendees (one per line)",
            height=140,
            placeholder="John Doe - (JD)",
        )
    with a2:
        att_ext = st.text_area(
            "🤝 External attendees (one per line)",
            height=140,
            placeholder="Jane Smith - (JS)",
        )

    if "preview_ready" not in st.session_state:
        st.session_state.preview_ready = False

    st.markdown("<div style='margin: 2rem 0;'></div>", unsafe_allow_html=True)

    # Preview builder (agent-driven)
    if st.button("🚀 Generate Weekly Preview", use_container_width=True, type="primary"):

        if project_tier == "T2" and not consolidate_week_with_agent_t2:
            st.error("Weekly T2 agent import failed. Ensure modules/nfr/nfr_weekly_agent_T2.py exists.")
            st.stop()
        if project_tier == "T1" and not consolidate_week_with_agent_t1:
            st.error("Weekly T1 agent import failed. Ensure modules/nfr/nfr_weekly_agent_T1.py exists.")
            st.stop()

        overview_agent = {
            "week_commencing": mon.strftime(DFMT),
            "date_range": date_range_label,
            "meeting_title": merged_config.get("meeting_title", "Weekly Catch Up"),
            "project_client": f"{project_name} — {client_name}",
        }

        provided_int = [x.strip() for x in (att_int or "").splitlines() if x.strip()]
        provided_ext = [x.strip() for x in (att_ext or "").splitlines() if x.strip()]

        # Extract actions from uploaded daily docs (optional)
        extracted_actions_rows = []
        for f in (weekly_files or []):
            try:
                d = Document(f)  # UploadedFile is file-like
                extracted_actions_rows.extend(find_actions_tables(d))
            except Exception:
                pass
        extracted_actions = consolidate_actions(extracted_actions_rows)

        # TODO (later): pull actual daily_nfr JSONs from DB and pass as daily_structured_inputs
        daily_structured_inputs = []

        with st.spinner(f"🔄 Consolidating week with agent ({project_tier})…"):
            if project_tier == "T1":
                out = consolidate_week_with_agent_t1(
                    overview=overview_agent,
                    monday_text=monday_text,
                    daily_structured_inputs=daily_structured_inputs,
                    extracted_actions=extracted_actions,
                    provided_attendees_internal=provided_int,
                    provided_attendees_external=provided_ext,
                )
            else:
                out = consolidate_week_with_agent_t2(
                    overview=overview_agent,
                    monday_text=monday_text,
                    daily_structured_inputs=daily_structured_inputs,
                    extracted_actions=extracted_actions,
                    provided_attendees_internal=provided_int,
                    provided_attendees_external=provided_ext,
                )

        weekly_structured = out.structured_data
        agent_flags = out.flags or []

        # Convert agent discussion schema -> tuples for existing preview/doc functions
        discussion_tuples = []
        for d in (weekly_structured.get("discussion") or []):
            discussion_tuples.append((d.get("subhead", ""), d.get("bullets", []) or []))

        st.session_state.preview = {
            "overview": {
                "Week commencing": overview_agent["week_commencing"],
                "Date range": overview_agent["date_range"],
                "Meeting title": overview_agent["meeting_title"],
                "Project / Client": overview_agent["project_client"],
            },
            "objectives": weekly_structured.get("objectives", "") or "",
            "att_int": weekly_structured.get("attendees_internal", []) or [],
            "att_ext": weekly_structured.get("attendees_external", []) or [],
            "discussion": discussion_tuples,
            "actions": weekly_structured.get("actions", []) or [],
            "weekly_structured": weekly_structured,
            "agent_flags": agent_flags,
            "project_tier": project_tier,
            "filename": f"Weekly_NFR_{mon.isoformat()}_{project_name.replace(' ', '_')}.docx",
            "week_commencing": mon,
            "project_name": project_name,
            "project_id": project_id,
            "client_name": client_name,
            "client_id": client_id,
            "client_code": client_code,
            "project_code": project_code,
        }
        st.session_state.preview_ready = True
        st.rerun()

    # Preview + Approve
    if st.session_state.get("preview_ready"):
        st.markdown(
            """
        <div class='section-header'>
            <h3>🔍 Weekly NFR Preview</h3>
        </div>
        """,
            unsafe_allow_html=True,
        )

        pv = st.session_state.preview

        st.markdown("<div class='preview-container'>", unsafe_allow_html=True)
        html = as_html_preview(
            pv["overview"],
            pv["objectives"],
            pv["att_int"],
            pv["att_ext"],
            pv["discussion"],
            pv["actions"],
        )
        st.components.v1.html(html, height=800, scrolling=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Agent flags (helps user spot missing/uncertain areas)
        if pv.get("agent_flags"):
            st.warning("⚠ Weekly agent flags (review before approving):")
            st.write(pv["agent_flags"])

        st.markdown(
            """
        <div class='info-box'>
            <strong style='color: #48bb78;'>📋 Review Instructions</strong><br/>
            <span style='color: #2d3748;'>Review the preview above. When you're satisfied with the content, click the button below to approve and generate the final document.</span>
        </div>
        """,
            unsafe_allow_html=True,
        )

        if st.button("✅ Approve & Create .docx", use_container_width=True, type="primary"):

            # Build DOCX
            data_bytes = make_weekly_docx(
                pv["overview"],
                pv["objectives"],
                pv["att_int"],
                pv["att_ext"],
                pv["discussion"],
                pv["actions"],
            )

            filename = pv["filename"]
            generated_by = (st.session_state.get("email", "system") or "").strip().lower()

            # Engine output (agent-driven) — aligns with your weekly_nfr table columns
            weekly_structured = pv.get("weekly_structured") or {}
            engine_output = {
                "overview": {
                    "week_commencing": pv["overview"]["Week commencing"],
                    "date_range": pv["overview"]["Date range"],
                    "meeting_title": pv["overview"]["Meeting title"],
                    "project_client": pv["overview"]["Project / Client"],
                },
                "objectives": pv["objectives"],
                "attendees_internal": pv["att_int"],
                "attendees_external": pv["att_ext"],
                "discussion_sections": [{"subhead": s, "bullets": b} for (s, b) in pv["discussion"]],
                "actions": pv["actions"],
                "agent_flags": pv.get("agent_flags", []),
                "agent_tier": pv.get("project_tier", "T2"),
                "raw_agent": weekly_structured,
            }

            # ---------------------------------------------------------
            # SAVE TO DB (weekly_nfr + weekly_nfr_files)
            # ---------------------------------------------------------
            weekly_nfr_id = None
            try:
                # insert_weekly_nfr is assumed to return weekly_nfr_id
                weekly_nfr_id = insert_weekly_nfr(
                    client_id=pv["client_id"],
                    project_id=pv["project_id"],
                    week_commencing=pv["week_commencing"],
                    date_range=pv["overview"]["Date range"],
                    meeting_title=pv["overview"]["Meeting title"],
                    project_client=pv["overview"]["Project / Client"],
                    objectives=pv["objectives"],
                    attendees_internal=json.dumps(pv["att_int"]),
                    attendees_external=json.dumps(pv["att_ext"]),
                    discussion_sections=json.dumps(engine_output["discussion_sections"]),
                    actions=json.dumps(pv["actions"]),
                    file_name=filename,
                    generated_by=generated_by,
                    raw_json=json.dumps(engine_output),
                )

                save_weekly_nfr_file(
                    weekly_nfr_id=weekly_nfr_id,
                    file_bytes=data_bytes,
                )

                # Log event (separate from notifications)
                log_event(
                    "weekly_nfr_created",
                    {
                        "user_email": generated_by,
                        "client_id": pv["client_id"],
                        "project_id": pv["project_id"],
                        "week_commencing": pv["week_commencing"].isoformat(),
                        "file_name": filename,
                        "tier": pv.get("project_tier", "T2"),
                        "weekly_nfr_id": int(weekly_nfr_id) if weekly_nfr_id is not None else None,
                    },
                )

                # ---------------------------------------------------------
                # NOTIFICATIONS (v5): Weekly NFR generated
                # ---------------------------------------------------------
                safe_emit_event(
                    event_type="nfr.weekly_generated",
                    project_id=pv["project_id"],
                    actor_email=generated_by,
                    title=f"Weekly NFR generated: {pv['overview']['Meeting title']}",
                    body=(
                        f"Client: {pv['client_name']}\n"
                        f"Project: {pv['project_name']}\n"
                        f"Week commencing: {pv['week_commencing'].strftime('%d %B %Y')}\n"
                        f"Weekly NFR ID: {weekly_nfr_id}\n"
                        f"File: {filename}"
                    ),
                    entity_type="weekly_nfr",
                    entity_id=int(weekly_nfr_id) if weekly_nfr_id is not None else None,
                )

            except Exception as e:
                st.error(f"❌ Could not save Weekly NFR to DB: {e}")
                st.stop()

            st.success("🎉 Weekly NFR generated and saved successfully!")

            # Download button
            st.download_button(
                "📄 Download Weekly NFR (.docx)",
                data_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary",
            )

            # Debug expander
            with st.expander("🔍 View Weekly Agent JSON", expanded=False):
                st.json(engine_output)

            st.info(f"💾 Saved to database with Weekly NFR ID: {weekly_nfr_id}")


# -----------------------------------------------------------
# FOOTER
# -----------------------------------------------------------
st.markdown("<div style='margin: 4rem 0 2rem 0;'></div>", unsafe_allow_html=True)
pmo_footer()

ui()
